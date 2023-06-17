import streamlit as st
import requests
import json
import io
import utils
from docx import Document
from docx.shared import Pt
import datetime
import zipfile
import re

utils.set_page_configuration(initial_sidebar_state="collapsed")
utils.set_sidebar()

# Pappers setup
url = "https://api.pappers.fr/v2/entreprise"
api_token = utils.open_file('pappers_api_key.txt')


title = 'Company sheet generator'
subtitle = "Automatically create the Company Sheet from a SIRET"
mission = "Connected to Pappers API, this app automatically generates the Company Sheet from a SIRET you provide. The only thing you have to do is to enter the SIRET of the company you are interested in."
st.markdown(f"<h1 style='text-align: center;font-size:50px;'>{title}</h1>", unsafe_allow_html=True)
st.markdown(f"<h3 style='text-align: center;'><i>{subtitle}</i></h3>", unsafe_allow_html=True)
st.markdown(f"<p style='text-align: justify;'>{mission}</p>", unsafe_allow_html=True)

st.text('')
siret = st.text_input('SIRET:', placeholder='Ex: 34300986600611', label_visibility="visible")
siret = siret.replace(' ', '').replace('-', '').replace('.', '')
st.text('')
st.text('')

cols = st.columns(3)
run = cols[1].button('Generate Company Sheet', key=None, help=None, on_click=None, args=None, kwargs=None, disabled=False)

if run:
    st.text('')

    params = {"api_token": api_token, "siret": siret}
    response = requests.get(url, params=params)

    if response.status_code == 200:
        data = response.json()
        # Process the retrieved data as needed
    
        # data = json.load(open('data/google.json', 'r', encoding='utf-8'))
        doc = Document('assets/template_fiche_societe.docx')

        # Create company name
        if data['nom_entreprise'] != None:
            company_name = data['nom_entreprise']
        else:
            company_name = 'Not specified ("nom_entreprise" missing)'

        # Create corporate form
        if data['forme_juridique'] != None:
            corporate_form = data['forme_juridique']
        else:
            corporate_form = 'Not specified ("forme_juridique" missing)'

        # Create company purpose
        if data['objet_social'] != None:
            company_purpose = data['objet_social']
        else:
            company_purpose = 'Not specified ("objet_social" missing)'

        # Create country
        if data['siege']['pays'] != None:
            country = data['siege']['pays']
        else:
            country = 'Not specified ("siège" > "pays" missing)'

        # Create address
        if data['siege']['adresse_ligne_1'] != None or data['siege']['code_postal'] != None or data['siege']['ville'] != None or data['siege']['pays'] != None:
            registered_office = data['siege']['adresse_ligne_1'] + ', ' + data['siege']['code_postal'] + ', ' + data['siege']['ville'] + ', ' + data['siege']['pays']
        else:
            registered_office = 'Not specified ("siège" > "adresse_ligne_1", "code_postal", "ville" or "pays" missing)'

        # Create Share Capital
        if data['capital_formate'] != None:
            share_capital = data['capital_formate']
        else:
            share_capital = 'Not specified ("capital_formate" missing)'

        # Create registration number
        if data['siren_formate'] != None:
            registration_number = data['siren_formate']
        else:
            registration_number = 'Not specified ("siren_formate" missing)'

        # Create RCS inscription
        if data['statut_rcs'] != None or data['greffe'] != None or data['date_immatriculation_rcs'] != None:
            rcs_inscription = f"{data['statut_rcs']} (au greffe de {data['greffe']}, le {data['date_immatriculation_rcs']})"
        else:
            rcs_inscription = 'Not specified ("statut_rcs", "greffe" or "date_immatriculation_rcs" missing)'

        # Create management
        if data['representants'] != None:
            management = ", ".join([f"{el['nom_complet']} ({el['qualite']})" for el in data['representants'] if 'Commissaire aux comptes' not in el['qualite']])
        else:
            management = 'Not specified ("representants" missing)'

        # Create statutory auditors
        if data['representants'] != None:
            if len([f"{el['nom_complet']} ({el['qualite']})" for el in data['representants'] if 'Commissaire aux comptes' in el['qualite']]) > 0:
                if data['representants']['siren'] != None:
                    statutory_auditors_principals = " ".join([f"{el['nom_complet']} ({el['siren']})" for el in data['representants'] if 'Commissaire aux comptes titulaire' in el['qualite']])
                    statutory_auditors_alternates = " ".join([f"{el['nom_complet']} ({el['siren']})" for el in data['representants'] if el['qualite'] == 'Commissaire aux comptes suppléant'])
                else:
                    statutory_auditors_principals = " ".join([f"{el['nom_complet']}" for el in data['representants'] if 'Commissaire aux comptes titulaire' in el['qualite']])
                    statutory_auditors_alternates = " ".join([f"{el['nom_complet']}" for el in data['representants'] if el['qualite'] == 'Commissaire aux comptes suppléant'])
            else:
                statutory_auditors_principals = 'No "Commissaire aux comptes" specified'
                statutory_auditors_alternates = 'No "Commissaire aux comptes" specified'
        else:
            statutory_auditors_principals = 'No "Commissaire aux comptes" specified'
            statutory_auditors_alternates = 'No "Commissaire aux comptes" specified'

        # Create fiscal year
        FRENCH_MONTHS = {
            'janvier': 'January',
            'février': 'February',
            'fevrier': 'February',
            'mars': 'March',
            'avril': 'April',
            'mai': 'May',
            'juin': 'June',
            'juillet': 'July',
            'août': 'August',
            'aout': 'August',
            'septembre': 'September',
            'octobre': 'October',
            'novembre': 'November',
            'décembre': 'December',
            'decembre': 'December',
        }
        
        def convert_fr_date_to_datetime(date_str):
            day, month = date_str.split()
            month = FRENCH_MONTHS[month.lower()]
            return datetime.datetime.strptime(f"{day} {month}", "%d %B")
        
        if data['date_cloture_exercice'] != None:
            end_date = convert_fr_date_to_datetime(data['date_cloture_exercice'])
            start_date = (end_date + datetime.timedelta(days=1)).strftime("%B, %d")
            end_date = end_date.strftime("%B, %d")
            fiscal_year = f'{start_date} to {end_date}'
            fiscal_year = fiscal_year.replace(' 0', ' ')
        else:
            fiscal_year = 'Not specified'
        
        # Create distribution of dividends
        lastest_exercices = [el['annee'] for el in data['finances']]
        lastest_exercices.sort(reverse=True)
        if len(lastest_exercices) >= 3:
            lastest_exercices = lastest_exercices[:3]
            distribution_of_dividends = {
                'year_1': str(lastest_exercices[0]),
                'net_income_1': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == lastest_exercices[0]][0]),
                'dividends_1': '', 'distribution_date_1': '',
                'year_2': str(lastest_exercices[1]),
                'net_income_2': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == lastest_exercices[1]][0]),
                'dividends_2': '', 'distribution_date_2': '',
                'year_3': str(lastest_exercices[2]),
                'net_income_3': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == lastest_exercices[2]][0]),
                'dividends_3': '', 'distribution_date_3': ''
            }
        elif len(lastest_exercices) == 2:
            two_lastest_exercices = lastest_exercices[:2]
            distribution_of_dividends = {
                'year_1': two_lastest_exercices[0],
                'net_income_1': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == two_lastest_exercices[0]][0]),
                'dividends_1': '', 'distribution_date_1': '',
                'year_2': two_lastest_exercices[1],
                'net_income_2': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == two_lastest_exercices[1]][0]),
                'dividends_2': '', 'distribution_date_2': '', 'year_3': '', 'net_income_3': '', 'dividends_3': '', 'distribution_date_3': ''
            }
        elif len(lastest_exercices) == 1:
            lastest_exercice = lastest_exercices[0]
            distribution_of_dividends = {
                'year_1': lastest_exercice,
                'net_income_1': str([el['chiffre_affaires'] for el in data['finances'] if el['annee'] == lastest_exercice][0]),
                'dividends_1': '', 'distribution_date_1': '', 'year_2': '', 'net_income_2': '', 'dividends_2': '', 
                'distribution_date_2': '', 'year_3': '', 'net_income_3': '', 'dividends_3': '', 'distribution_date_3': ''
            }
        else:
            distribution_of_dividends = {
                'year_1': '', 'net_income_1': '', 'dividends_1': '', 'distribution_date_1': '', 'year_2': '', 'net_income_2': '', 'dividends_2': '', 
                'distribution_date_2': '', 'year_3': '', 'net_income_3': '', 'dividends_3': '', 'distribution_date_3': ''
            }

        # Create term
        if data['duree_personne_morale'] != None and data['date_immatriculation_rcs'] != None:
            term = (datetime.datetime.strptime(data['date_immatriculation_rcs'], '%Y-%m-%d') + datetime.timedelta(days=data['duree_personne_morale']*365)).strftime('%Y-%m-%d')
        else:
            term = 'Not specified ("date_immatriculation_rcs" or "duree_personne_morale" missing))'

        replacements = {
            'country': country,
            'company_name': company_name,
            'corporate_form': corporate_form,
            'registered_office': registered_office,
            'share_capital': share_capital,
            # 'issued_shares': 'Not provided',
            # 'shareholding': 'Not provided',
            'registration_number': registration_number,
            'rcs_inscription': rcs_inscription,
            'company_purpose': company_purpose,
            'term': term,
            'fiscal_year': fiscal_year,
            # 'management_president': [x for x in data['representants'] if x['qualite'] == 'Président'][0]['nom_complet'],
            # 'management_board_of_directors': [x['nom_complet'] for x in data['representants'] if x['qualite'] == 'Gérant'],
            'management': management,
            'statutory_auditors_principals': statutory_auditors_principals,
            'statutory_auditors_alternates': statutory_auditors_alternates,
            # 'statutory_restrictions_on_transfer_of_shares': 'Not provided',
            # 'powers_restriction_board_of_directors_reserved_matters': 'Not provided',
            # 'other_securities_giving_access_to_the_share_capital': 'Not provided',
            # 'pledge_over_securities': 'Not provided',
            'year1': distribution_of_dividends['year_1'],
            'net_income1': distribution_of_dividends['net_income_1'],
            'dividends1': distribution_of_dividends['dividends_1'],
            'distributiondate1': distribution_of_dividends['distribution_date_1'],
            'year2': distribution_of_dividends['year_2'],
            'net_income2': distribution_of_dividends['net_income_2'],
            'dividends2': distribution_of_dividends['dividends_2'],
            'distributiondate2': distribution_of_dividends['distribution_date_2'],
            'year3': distribution_of_dividends['year_3'],
            'net_income3': distribution_of_dividends['net_income_3'],
            'dividends3': distribution_of_dividends['dividends_3'],
            'distributiondate3': distribution_of_dividends['distribution_date_3'],
            'date_of_updated_articles_of_association': data['derniere_mise_a_jour_sirene']
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            run.font.name = 'Arial'
                            run.font.size = Pt(8)
                            run.font.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                run = paragraph.runs
                                for i in range(len(run)):
                                    if key in run[i].text:
                                        run[i].text = run[i].text.replace(key, value)
                                        run[i].font.name = 'Arial'
                                        run[i].font.size = Pt(8)

        # doc.save('results/filled_document.docx')
        # bio = io.BytesIO()
        # doc.save(bio)

        # st.text('')
        # st.text('')
        # st.markdown(f"<h3 style='text-align: center;'>The Company Sheet is ready (already) 🎁</h3>", unsafe_allow_html=True)
        # st.text('')
        # cols = st.columns(3)
        # cols[1].download_button(
        #     label="Download Company Sheet",
        #     data=bio.getvalue(),
        #     file_name=f"{replacements['company_name'].replace(' ', '_').lower()}_sheet.docx",
        #     mime="docx",
        #     key='docx')

        bio = io.BytesIO()
        doc.save(bio)
        docx_data = bio.getvalue()

        # Save json
        bio = io.BytesIO()
        bio.write(json.dumps(data, indent=4).encode('utf-8'))
        json_data = bio.getvalue()

        def remove_parentheses_content(s):
            return re.sub("\([^)]*\)", '', s)

        # Create a zip file
        bio = io.BytesIO()
        with zipfile.ZipFile(bio, 'w') as zipf:
            zipf.writestr(f"{remove_parentheses_content(replacements['company_name'].replace(' ', '_').lower())}_sheet.docx", docx_data)
            zipf.writestr(f"{remove_parentheses_content(replacements['company_name'].replace(' ', '_').lower())}_api_response.json", json_data)

        st.text('')
        st.text('')
        st.markdown(f"<h3 style='text-align: center;'>The Company Sheet is ready 🎁</h3>", unsafe_allow_html=True)
        st.text('')
        cols = st.columns(3)
        cols[1].download_button(
            label="Download Company Sheet",
            data=bio.getvalue(),
            file_name=f"{remove_parentheses_content(replacements['company_name'].replace(' ', '_').lower())}.zip",
            mime="application/zip",
            key='zip')
    
    elif response.status_code == 400:
        st.write(f"⚠️ The SIRET number {siret} is not valid. Please check the number and try again.")
        st.write(f"Make sure you enter the 14-digit SIRET number, not the 9-digit SIREN number.")

    else:
        st.write(f"Error occurred: {response.status_code}")    